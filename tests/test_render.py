"""Renderers: PDFs must be exactly one page; DOCX must be valid and complete."""
from __future__ import annotations

import re

import pytest
from docx import Document

from app.render_docx import render_cover_letter_docx, render_resume_docx
from app.render_pdf import (
    render_cover_letter_pdf,
    render_resume_pdf,
    resume_fit_scale,
)
from app.schema import Contact, ExperienceItem, ResumeContent, SkillLine


def _overflowing_resume() -> ResumeContent:
    """A resume far over the one-page budget, so the auto-fit must shrink type."""
    return ResumeContent(
        contact=Contact(name="Overflow Candidate", email="o@e.com", phone="+1-555",
                        links=["linkedin.com/in/overflow", "github.com/overflow"]),
        summary="A deliberately long professional summary " * 6,
        skills=[SkillLine(label=f"Skill area {i}", items="Tool A, Tool B, Tool C, Tool D")
                for i in range(6)],
        experience=[ExperienceItem(
            company=f"Company {i}", title="Senior Staff Engineer", dates="2018 - 2024",
            bullets=[f"Delivered measurable outcome {i}-{j} with a clear quantified metric."
                     for j in range(5)],
        ) for i in range(10)],
        education=["B.Tech - IIT Madras, 2011", "M.S. - Example University, 2014"],
    )


def _name_run_size_pt(path: str, name: str) -> float:
    for p in Document(path).paragraphs:
        if p.text == name and p.runs:
            return p.runs[0].font.size.pt
    raise AssertionError(f"name paragraph {name!r} not found")


def _pdf_page_count(data: bytes) -> int:
    """Count page objects in a reportlab PDF.

    reportlab writes each page as an uncompressed `/Type /Page` object and one
    `/Type /Pages` root; the `\\b` keeps `/Pages` from matching `/Page`.
    """
    return len(re.findall(rb"/Type\s*/Page\b", data))


def _docx_text(path: str) -> str:
    return "\n".join(p.text for p in Document(path).paragraphs)


def test_resume_pdf_is_one_page(docs, tmp_path):
    out = tmp_path / "resume.pdf"
    render_resume_pdf(docs.resume, str(out))
    data = out.read_bytes()
    assert data.startswith(b"%PDF")
    assert _pdf_page_count(data) == 1


def test_dense_resume_still_one_page(dense_docs, tmp_path):
    """Auto-fit must shrink a near-budget resume down to a single page."""
    out = tmp_path / "dense.pdf"
    render_resume_pdf(dense_docs.resume, str(out))
    assert _pdf_page_count(out.read_bytes()) == 1


def test_cover_letter_pdf_is_one_page(docs, tmp_path):
    out = tmp_path / "cover.pdf"
    render_cover_letter_pdf(docs.cover_letter, docs.resume.contact, str(out))
    data = out.read_bytes()
    assert data.startswith(b"%PDF")
    assert _pdf_page_count(data) == 1


def test_resume_docx_is_valid_and_complete(docs, tmp_path):
    out = tmp_path / "resume.docx"
    render_resume_docx(docs.resume, str(out))
    text = _docx_text(str(out))
    assert "Rashmi Mahadevaiah" in text
    assert "EXPERIENCE" in text          # headings are upper-cased
    assert "Languages & ML:" in text     # skill label rendered with colon
    assert "Built a RAG pipeline serving 2M requests/day." in text


def test_cover_letter_docx_is_valid_and_complete(docs, tmp_path):
    out = tmp_path / "cover.docx"
    render_cover_letter_docx(docs.cover_letter, docs.resume.contact, str(out))
    text = _docx_text(str(out))
    assert "Rashmi Mahadevaiah" in text
    assert "Dear Hiring Manager," in text
    assert "Sincerely," in text


def test_cover_letter_docx_falls_back_to_contact_name(docs, tmp_path):
    docs.cover_letter.signature = ""
    out = tmp_path / "cover.docx"
    render_cover_letter_docx(docs.cover_letter, docs.resume.contact, str(out))
    # signature line falls back to the contact name
    assert "Rashmi Mahadevaiah" in _docx_text(str(out))


def test_resume_fit_scale_full_when_fits(docs):
    """A normal-length resume fits at full size — no shrink needed."""
    assert resume_fit_scale(docs.resume) == 1.0


def test_resume_fit_scale_shrinks_when_overflowing():
    """A resume far over the one-page budget forces the fit below full size."""
    assert resume_fit_scale(_overflowing_resume()) < 1.0


def test_resume_docx_mirrors_scale(docs, tmp_path):
    """The DOCX type scales with the fit factor, so Word matches the PDF.

    OOXML stores font size in half-points, so 19*0.8=15.2pt reads back as 15.0pt;
    compare approximately rather than exactly.
    """
    name = docs.resume.contact.name
    full = tmp_path / "full.docx"
    shrunk = tmp_path / "shrunk.docx"
    render_resume_docx(docs.resume, str(full), scale=1.0)
    render_resume_docx(docs.resume, str(shrunk), scale=0.8)
    full_pt = _name_run_size_pt(str(full), name)
    shrunk_pt = _name_run_size_pt(str(shrunk), name)
    assert full_pt == pytest.approx(19, abs=0.3)
    assert shrunk_pt == pytest.approx(19 * 0.8, abs=0.3)
    assert shrunk_pt < full_pt


def test_overflowing_resume_docx_uses_shrunk_scale(tmp_path):
    """End to end: an overflowing resume's DOCX renders at the shrunk fit scale,
    so its type is smaller than a full-size render of the same content."""
    resume = _overflowing_resume()
    scale = resume_fit_scale(resume)
    assert scale < 1.0
    full = tmp_path / "full.docx"
    fitted = tmp_path / "fitted.docx"
    render_resume_docx(resume, str(full), scale=1.0)
    render_resume_docx(resume, str(fitted), scale=scale)
    name = resume.contact.name
    assert _name_run_size_pt(str(fitted), name) < _name_run_size_pt(str(full), name)
