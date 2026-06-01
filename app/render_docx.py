"""Render TailoredDocs to Word (.docx) with python-docx.

Layout matches the PDF renderer: navy name, ruled headings, labelled skills,
role rows with right-tabbed dates, bullet lists. python-docx has no layout
engine, so one-page fit is achieved through compact typography and the prompt's
length budgets rather than a measured auto-fit (see docs/TECH_DEBT.md).
"""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Twips

from .schema import Contact, CoverLetter, ResumeContent

NAVY = RGBColor(0x1F, 0x3B, 0x5B)
RULE = "2E5A88"
GRAY = RGBColor(0x55, 0x55, 0x55)
CONTENT_W_TWIPS = 9746  # A4 width 11906 - left 1080 - right 1080
FONT = "Arial"


def _runs(paragraph, scale, *runs):
    """Append (text, size_pt, bold, color) runs, scaling each size by `scale`."""
    for text, size, bold, color in runs:
        r = paragraph.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size * scale)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
    return paragraph


def _add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")     # eighths of a point => 0.75pt
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), RULE)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _heading(doc, text, scale):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 * scale)
    p.paragraph_format.space_after = Pt(4 * scale)
    _runs(p, scale, (text.upper(), 12, True, NAVY))
    _add_bottom_border(p)
    return p


def _contact_line(c: Contact) -> str:
    parts = [c.email, c.phone, *c.links]
    return "   |   ".join(p for p in parts if p)


def _base_document(scale: float = 1.0) -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10 * scale)
    # Pin spacing so only the explicit per-paragraph values count; Word's default
    # template otherwise adds ~8pt after every paragraph and 1.08 line spacing,
    # which silently pushes a tight resume onto a second page.
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Twips(720)
    sec.bottom_margin = Twips(620)
    sec.left_margin = Twips(1080)
    sec.right_margin = Twips(1080)
    return doc


def render_resume_docx(resume: ResumeContent, out_path: str, scale: float = 1.0) -> str:
    """Render the resume to .docx. `scale` (≤1) mirrors the PDF's one-page fit."""
    doc = _base_document(scale)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2 * scale)
    _runs(name_p, scale, (resume.contact.name, 19, True, NAVY))

    contact = _contact_line(resume.contact)
    if contact:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(2 * scale)
        _runs(cp, scale, (contact, 9, False, GRAY))

    if resume.summary:
        _heading(doc, "Summary", scale)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1 * scale)
        _runs(p, scale, (resume.summary, 10, False, None))

    if resume.skills:
        _heading(doc, "Technical Skills", scale)
        for s in resume.skills:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2 * scale)
            _runs(p, scale, (f"{s.label}:  ", 10, True, None), (s.items, 10, False, None))

    if resume.experience:
        _heading(doc, "Experience", scale)
        for e in resume.experience:
            rp = doc.add_paragraph()
            rp.paragraph_format.space_before = Pt(7 * scale)
            rp.paragraph_format.space_after = Pt(2 * scale)
            rp.paragraph_format.tab_stops.add_tab_stop(Twips(CONTENT_W_TWIPS), WD_TAB_ALIGNMENT.RIGHT)
            title = f"  —  {e.title}" if e.title else ""
            _runs(rp, scale,
                  (e.company, 11, True, None),
                  (title, 11, False, None))
            if e.dates:
                _runs(rp, scale, ("\t" + e.dates, 10, False, GRAY))
            for b in e.bullets:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(1.5 * scale)
                _runs(bp, scale, (b, 10, False, None))

    if resume.education:
        _heading(doc, "Education", scale)
        for line in resume.education:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1 * scale)
            _runs(p, scale, (line, 10, False, None))

    doc.save(out_path)
    return out_path


def render_cover_letter_docx(cl: CoverLetter, contact: Contact, out_path: str,
                             scale: float = 1.0) -> str:
    """Render the cover letter to .docx. `scale` (≤1) mirrors the PDF's one-page fit."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11 * scale)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Twips(1440)
    sec.bottom_margin = Twips(1080)
    sec.left_margin = Twips(1440)
    sec.right_margin = Twips(1440)

    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(1 * scale)
    _runs(name_p, scale, (contact.name, 16, True, NAVY))

    line = _contact_line(contact)
    if line:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_after = Pt(14 * scale)
        _runs(cp, scale, (line, 9, False, GRAY))

    def meta(text, after=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after * scale)
        _runs(p, scale, (text, 11, False, None))
        return p

    if cl.date:
        meta(cl.date, after=10)
    for r in cl.recipient:
        if r:
            meta(r, after=0)
    if cl.recipient:
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    if cl.salutation:
        meta(cl.salutation, after=8)

    for para in cl.paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10 * scale)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _runs(p, scale, (para, 11, False, None))

    if cl.closing:
        meta(cl.closing, after=2)
    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(10 * scale)
    _runs(sig, scale, (cl.signature or contact.name, 11, True, None))

    doc.save(out_path)
    return out_path
